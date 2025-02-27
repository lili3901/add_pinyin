from docx import Document
from docx.shared import Pt  # 用于设置字体大小
from pypinyin import pinyin, Style
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

# 指定要处理的Word文档路径
doc_path = 'C:/Users/lilil/Desktop/G2P2.docx'
output_path = 'C:/Users/lilil/Desktop/annotated_G2P2.docx'

# 打开现有的Word文档
try:
    doc = Document(doc_path)
except Exception as e:
    print(f"无法打开文档: {e}")
    exit(1)

# 创建一个新的文档来保存结果
new_doc = Document()

# 定义每行的最大列数
max_columns_per_row = 15
max_width=230
normal_width=9

# 遍历文档中的所有段落
for paragraph in doc.paragraphs:
    # 添加一个空行以分隔不同的表格块
    new_doc.add_paragraph()
    text = paragraph.text.strip()    
    # if not text:  # 如果是空段落则跳过
    #     new_doc.add_paragraph()  # 保持原始文档中的空行
    #     continue
    # 若没有应用了“正文”样式，则复制样式
    if paragraph.style.name != 'Normal':      
        current_para=new_doc.add_paragraph(text)
        current_para.style = paragraph.style.name
        continue
    # 初始化变量
    current_table = None
    column_count = 0
    is_new_paragraph=True
    # 获取拼音列表（带声调符号），并处理每个字符
    # data用于临时存放行数据
    data=[[],[],[]]
    total_width=0
    total_chars=len(text)
    current_index=0
    for idx, char in enumerate(text):
        current_index += 1
        # 当达到每行宽度后才开始绘制表格
            # 在新段落前空两格
        if is_new_paragraph:
            is_new_paragraph=False
            for i in range(2): 
                for row in range(2):            
                    data[row].append("")
                data[2].append(normal_width)
                total_width += normal_width
                column_count += 1
        if '\u4e00' <= char <= '\u9fff':  # 判断是否为汉字
            # 获取单个汉字的拼音列表（可能有多个读音）
            pinyin_list = pinyin(char, style=Style.TONE)
            # 取第一个拼音（对于多音字，这里只取最常见的一种）
            pinyin_str = pinyin_list[0][0] if pinyin_list else ''
            
            # 添加拼音到第一行
            data[0].append(pinyin_str)
            cell_width=max(len(pinyin_str)*5.3,8)
            total_width += cell_width
            data[2].append(cell_width)
            # 添加汉字到第二行
            data[1].append(char)
            column_count += 1

        else:
            # 如果不是汉字，则在汉字行中都添加该字符
            data[0].append("")
            data[1].append(char)
            data[2].append(normal_width)
            total_width += normal_width
            column_count += 1



        if total_width >= max_width or current_index==total_chars:
            # 当达到最大宽度开始时创建新表格
            
            # 添加一个空行以分隔不同的表格块
            paragraph=new_doc.add_paragraph()
            paragraph_format=paragraph.paragraph_format
            paragraph_format.line_spacing=0.1
            paragraph_format.space_before=Pt(0)
            paragraph_format.space_after=Pt(0)
            min_cols=max(15,column_count)
            current_table = new_doc.add_table(rows=2, cols=min_cols)
           
            for cell in current_table._cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            # 设置第一行的高度,这里设置了一个过小的值
            current_table.rows[0].height = 104400  # 行高单位是EMU（English Metric Unit），1 inch = 914400 EMU
            current_table.rows[0].height_rule = None  # 或者使用 WD_ROW_HEIGHT_RULE.EXACTLY 来精确指定高度
            
            # 设置表格样式和宽度
            # current_table.style = 'Table Grid'  # 使用默认网格样式，可根据需要更改
            # for row in current_table.rows:
            #     for cell in row.cells:
            #         cell.width = Inches(0.6)  # 设置单元格宽度
            #         for paragraph in cell.paragraphs:
            #             paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中文本

            for col in range(column_count):
                #current_table.cell(0,col).text=data[0][col]
                cell_pinyin = current_table.cell(0, col)
                cell_pinyin.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                cell_pinyin.width=Pt(data[2][col])
                run = cell_pinyin.paragraphs[0].add_run(data[0][col])
                run.font.size = Pt(8)  # 设置拼音字体大小                
                #current_table.cell(1,col).text=data[1][col]
                cell_hanzi = current_table.cell(1, col)
                cell_hanzi.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                run = cell_hanzi.paragraphs[0].add_run(data[1][col])
                run.font.name = 'KaiTi'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                run.font.size = Pt(13)  # 设置汉字字体大小

                #current_table.columns[col].width=Pt(data[2][col])
            column_count = 0
            data=[[],[],[]]
            total_width=0

        
"""         if '\u4e00' <= char <= '\u9fff':  # 判断是否为汉字
            # 获取单个汉字的拼音列表（可能有多个读音）
            pinyin_list = pinyin(char, style=Style.TONE)
            # 取第一个拼音（对于多音字，这里只取最常见的一种）
            pinyin_str = pinyin_list[0][0] if pinyin_list else ''
            
            # 添加拼音到第一行
            cell_pinyin = current_table.cell(0, column_count)
            cell_pinyin.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
            run = cell_pinyin.paragraphs[0].add_run(pinyin_str)
            run.font.size = Pt(8)  # 设置拼音字体大小
            
            # 添加汉字到第二行
            cell_hanzi = current_table.cell(1, column_count)
            cell_hanzi.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            run = cell_hanzi.paragraphs[0].add_run(char)
            run.font.name = 'KaiTi'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            run.font.size = Pt(13)  # 设置汉字字体大小
        else:
            # 如果不是汉字，则直接在两行中都添加该字符
            for row in range(2):
                cell = current_table.cell(row, column_count)
                if row==0:
                    run = cell.paragraphs[0].add_run("")                    
                else:
                    run = cell.paragraphs[0].add_run(char)
                    run.font.name = 'KaiTi'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        
        column_count += 1 """

# 保存新的文档到指定路径
try:
    new_doc.save(output_path)
    print(f"文档已成功保存至: {output_path}")
except Exception as e:
    print(f"无法保存文档: {e}")
