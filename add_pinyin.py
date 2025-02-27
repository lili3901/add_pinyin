from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from pypinyin import pinyin, Style

def add_pinyin_to_doc(input_path, output_path):
    doc = Document(input_path)

    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        original_runs = [{'text': run.text, 'font': run.font} for run in para.runs]
        para.clear()

        for run_info in original_runs:
            text = run_info['text']
            font = run_info['font']

            for char in text:
                # 生成拼音（过滤非汉字）
                if not '\u4e00' <= char <= '\u9fff':
                    para.add_run(char)  # 非汉字直接添加
                    continue

                pinyin_list = pinyin(char, style=Style.TONE)
                char_pinyin = pinyin_list[0][0] if pinyin_list else ''

                # 创建新Run并设置字体
                new_run = para.add_run()
                new_run.font.name = font.name
                new_run.font.size = font.size
                
                # 确保中文字体设置
                rPr = new_run._element.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:eastAsia'), font.name or "SimSun")

                # 构建Ruby标签
                ruby = OxmlElement('w:ruby')
                
                # 1. 拼音部分 (rt)
                rt = OxmlElement('w:rt')
                rt.text = char_pinyin
                
                # 手动添加拼音格式设置
                rPr_rt = OxmlElement('w:rPr')         # 新建格式属性
                sz_rt = OxmlElement('w:sz')           # 字号标签
                sz_rt.set(qn('w:val'), "20")          # 字号值（20=10磅）
                rPr_rt.append(sz_rt)
                rt.insert(0, rPr_rt)                  # 将格式添加到rt开头
                
                # 2. 基础文本部分 (rubyBase)
                ruby_base = OxmlElement('w:rubyBase')
                ruby_base.text = char
                
                # 3. 组装Ruby标签
                ruby.append(rt)
                ruby.append(ruby_base)
                
                # 4. 将Ruby添加到Run
                new_run._element.append(ruby)

    doc.save(output_path)

# 指定要处理的Word文档路径

input_doc = 'C:/Users/lilil/Desktop/yqzddwj.docx'
output_doc = 'C:/Users/lilil/Desktop/annotated_yqzddwj.docx'
add_pinyin_to_doc(input_doc, output_doc)